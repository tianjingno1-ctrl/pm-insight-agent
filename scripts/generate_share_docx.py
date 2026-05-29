"""Generate Word document for AI learning share presentation."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ai-learning-share.docx"


def set_doc_font(doc: Document, font_name: str = "微软雅黑") -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_meta(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.size = Pt(10)


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)


def build() -> None:
    doc = Document()
    set_doc_font(doc)

    add_title(doc, "PM Insight Agent")
    add_title(doc, "AI 学习实践分享")
    add_meta(
        doc,
        [
            "分享场景：需求评审会 · AI 学习心得",
            "分享人：[填写姓名]",
            "建议时长：15～20 分钟（含 5 分钟现场演示）",
            "项目：pm-insight-agent（分支 experiment/pony-roundtable-ui）",
            "文档版本：2026-05-29（演示 UI + 发言记录版）",
        ],
    )
    doc.add_page_break()

    add_h1(doc, "一、开场")
    add_h2(doc, "1.1 项目一句话")
    add_p(
        doc,
        "PM Insight Agent：面向产品经理的 AI 专家圆桌——你提一个议题，"
        "多个 AI 专家角色围绕它讨论，最后输出结构化小结。",
        bold=True,
    )
    add_h2(doc, "1.2 分享目标")
    add_bullets(
        doc,
        [
            "看见一种可能：AI 可以做成「多角色协作」，不只是聊天框",
            "理解边界：哪些是 Demo，哪些是真实工程能力",
            "可复用的方法：mock 先行、分阶段交付、Cursor 辅助开发",
        ],
    )

    add_h1(doc, "二、产品演进")
    add_bullets(
        doc,
        [
            "Phase 1 — CLI：粘贴材料 → 5 Agent → Markdown 报告 / PRD",
            "Phase 1.5 — Streamlit：多轮追问、三行小结、长期记忆",
            "Phase 2.1 — 圆桌 UI：四专家围坐、动画播放",
            "Phase 2.2 — SSE 事件流：Mock 先跑通协议",
            "Phase 2.3 — DeepSeek：真实 LLM 按议题生成讨论",
            "当前 — 演示 UI：发言记录汇总 + 短气泡 + 三列决策小结",
        ],
    )

    add_h1(doc, "三、界面说明（演示时指给观众看）")
    add_table(
        doc,
        ["区域", "作用"],
        [
            ["圆桌区", "四专家围坐；头顶短气泡显示互动，如「反驳产品」"],
            ["左侧发言记录", "全部专家发言汇总；当前发言流式更新；自动滚到底"],
            ["顶部状态徽章", "就绪 → LIVE 讨论进行中 → 讨论结束"],
            ["本轮决策小结", "三列卡片：当前倾向 / 最大分歧 / 下一步"],
        ],
    )

    add_h1(doc, "四、现场演示（5 分钟）")
    add_h2(doc, "4.1 启动")
    add_bullets(
        doc,
        [
            "终端 1：cd backend → py -3.12 -m uvicorn app.main:app --reload --port 8000",
            "终端 2：cd frontend → npm run dev",
            "浏览器：http://localhost:3000，按 F11 全屏",
        ],
    )
    add_h2(doc, "4.2 推荐议题")
    add_table(
        doc,
        ["议题", "说明"],
        [
            ["需求评审会要不要引入 AI 辅助？", "与分享主题直接相关（推荐）"],
            ["两周内 AI 财务助手 MVP 应该做什么？", "MVP 边界"],
            ["AI 会不会取代产品经理？", "多角色观点碰撞"],
        ],
    )
    add_h2(doc, "4.3 演示步骤")
    add_bullets(
        doc,
        [
            "输入议题 → 点击「开始讨论」",
            "等待 10～30 秒（DeepSeek 生成）",
            "指左侧发言记录 + 圆桌短气泡",
            "指底部三列决策小结",
        ],
    )
    add_h2(doc, "4.4 演示话术")
    add_quote(
        doc,
        "圆桌看互动关系，左侧看完整发言记录，最后看结构化小结。"
        "DeepSeek 驱动，失败自动降级，演示不会中断。",
    )
    add_h2(doc, "4.5 防翻车")
    add_table(
        doc,
        ["情况", "处理"],
        [
            ["API 失败", "自动 fallback，界面照常播"],
            ["页面卡住", "刷新重试"],
            ["端口占用", "重启前后端"],
        ],
    )

    add_h1(doc, "五、如何分享给别人看 ★")
    add_h2(doc, "5.1 现场演示（最推荐）")
    add_bullets(
        doc,
        [
            "你电脑启动项目，浏览器 F11 全屏",
            "接投影仪，或腾讯会议/飞书「共享屏幕」",
            "由你操作，观众观看——无需观众安装任何软件",
        ],
    )
    add_h2(doc, "5.2 发文档")
    add_table(
        doc,
        ["文件", "路径", "打开方式"],
        [
            ["Word 分享稿", "docs/ai-learning-share.docx", "Word / WPS 双击打开"],
            ["PPT 大纲", "docs/ai-learning-share-ppt-outline.md", "复制到 PowerPoint"],
        ],
    )
    add_p(doc, "通过微信、飞书、邮件发送 .docx 附件即可。")
    add_h2(doc, "5.3 录屏")
    add_bullets(
        doc,
        [
            "Windows：Win + G 打开录屏",
            "录完整流程：输入议题 → 讨论 → 决策小结",
            "导出 MP4 发群或上传内网",
        ],
    )
    add_h2(doc, "5.4 常见误区")
    add_table(
        doc,
        ["误区", "正确做法"],
        [
            ["把 localhost 链接发给别人", "localhost 只有你自己能开；请投屏或录屏"],
            ["让同事自己连你的 WiFi 访问", "当前后端绑本机，请由你演示"],
            ["未启动后端", "LLM 模式需 backend :8000"],
        ],
    )

    add_h1(doc, "六、AI 学习心得")
    add_bullets(
        doc,
        [
            "先 mock，再接真 API",
            "协议比模型更重要（MeetingEvent）",
            "多 Agent 分演示型 vs 工程型",
            "分阶段交付 + Git 锚点 + 交接文档",
            "演示稳定性 > 技术炫技",
        ],
    )

    add_h1(doc, "七、诚实边界")
    add_table(
        doc,
        ["已经做到", "还没做到"],
        [
            ["真实 LLM + 可视化圆桌 + 发言记录", "真正多 Agent 编排"],
            ["API 失败自动降级", "边生成边播（token 流式）"],
            ["演示级 UI", "生产部署 / RAG / 飞书集成"],
        ],
    )

    add_h1(doc, "八、Q&A 备答")
    qa = [
        ("套壳吗？", "多角色圆桌 + 事件协议，不是单轮问答。"),
        ("安全吗？", "本地原型；落地走公司模型网关。"),
        ("成本？", "DeepSeek 一次几分钱量级。"),
    ]
    for q, a in qa:
        add_p(doc, f"Q：{q}", bold=True)
        add_p(doc, f"A：{a}")

    add_h1(doc, "九、演示检查清单")
    add_bullets(
        doc,
        [
            "后端 /health 返回 200",
            "前端 localhost:3000 可打开",
            "DeepSeek Key 有效（或接受 fallback）",
            "议题已准备，F11 全屏试一遍",
            "（线上会）屏幕共享已测试",
        ],
    )

    add_h1(doc, "十、收尾金句")
    add_quote(
        doc,
        "AI 不会取代产品经理，但会用 AI 的产品经理会取代不会用的。"
        "目标是跑通「需求 → 多视角讨论 → 结构化输出」的 AI 链路。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Generated: {OUT}")


if __name__ == "__main__":
    build()
